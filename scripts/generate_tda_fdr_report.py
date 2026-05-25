from __future__ import annotations

from pathlib import Path
from typing import Iterable

import numpy as np
import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(r"C:\Users\tasne\Downloads")
OUT_DIR = Path(r"C:\Users\tasne\Desktop\project\outputs\tda_report")


def prefer_latest_csv(name: str) -> Path:
    numbered = BASE_DIR / f"{name} (1).csv"
    plain = BASE_DIR / f"{name}.csv"
    return numbered if numbered.exists() else plain


SUMMARY_CSV = prefer_latest_csv("tda_summary")
THRESHOLDS_CSV = prefer_latest_csv("tda_thresholds")
PSMS_CSV = prefer_latest_csv("tda_psms")


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 100, 112)
TEXT = RGBColor(31, 41, 55)
GRID = "D9E2F3"
HEADER_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
TARGET = "#2563EB"
DECOY = "#DC2626"


def font(size: int = 22, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except OSError:
            pass
    return ImageFont.load_default()


def fmt_num(value: float) -> str:
    value = float(value)
    if abs(value) >= 1000:
        return f"{value:,.0f}"
    if abs(value) >= 10:
        return f"{value:.1f}"
    if abs(value) >= 1:
        return f"{value:.2f}"
    out = f"{value:.3f}".rstrip("0").rstrip(".")
    return out if out else "0"


def nice_ticks(lo: float, hi: float, n: int = 6) -> list[float]:
    if hi <= lo:
        return [lo]
    return [lo + (hi - lo) * i / (n - 1) for i in range(n)]


def draw_histogram(
    values: np.ndarray,
    *,
    title: str,
    xlabel: str,
    color: str,
    path: Path,
    bins: int = 80,
    xlim: tuple[float, float] | None = None,
) -> None:
    values = values[np.isfinite(values)]
    width, height = 1800, 1050
    margin_left, margin_right, margin_top, margin_bottom = 150, 75, 130, 145
    plot_w = width - margin_left - margin_right
    plot_h = height - margin_top - margin_bottom
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)

    if xlim is None:
        lo, hi = float(values.min()), float(values.max())
        pad = (hi - lo) * 0.03 or 0.01
        lo, hi = lo - pad, hi + pad
    else:
        lo, hi = xlim

    counts, edges = np.histogram(values, bins=bins, range=(lo, hi))
    ymax = max(int(counts.max()), 1)
    ytop = ymax * 1.08

    def xmap(x: float) -> float:
        return margin_left + (x - lo) / (hi - lo) * plot_w

    def ymap(y: float) -> float:
        return margin_top + plot_h - y / ytop * plot_h

    # Background and grid.
    draw.rectangle(
        [margin_left, margin_top, margin_left + plot_w, margin_top + plot_h],
        fill="#F9FAFB",
        outline="#E5E7EB",
        width=2,
    )
    for tick in nice_ticks(0, ytop, 6):
        y = ymap(tick)
        draw.line([margin_left, y, margin_left + plot_w, y], fill="#E5E7EB", width=2)
        draw.text(
            (margin_left - 18, y),
            fmt_num(tick),
            fill="#4B5563",
            font=font(24),
            anchor="rm",
        )

    # Bars.
    for count, a, b in zip(counts, edges[:-1], edges[1:]):
        if count <= 0:
            continue
        x0, x1 = xmap(float(a)), xmap(float(b))
        y0 = ymap(float(count))
        draw.rectangle(
            [x0, y0, max(x0 + 1, x1 - 1), margin_top + plot_h],
            fill=color,
            outline="white",
            width=1,
        )

    # Axes and ticks.
    axis_color = "#111827"
    draw.line([margin_left, margin_top + plot_h, margin_left + plot_w, margin_top + plot_h], fill=axis_color, width=3)
    draw.line([margin_left, margin_top, margin_left, margin_top + plot_h], fill=axis_color, width=3)
    for tick in nice_ticks(lo, hi, 7):
        x = xmap(tick)
        draw.line([x, margin_top + plot_h, x, margin_top + plot_h + 11], fill=axis_color, width=2)
        draw.text((x, margin_top + plot_h + 32), fmt_num(tick), fill="#4B5563", font=font(24), anchor="mt")

    draw.text((width / 2, 58), title, fill="#111827", font=font(44, bold=True), anchor="mm")
    draw.text((margin_left + plot_w / 2, height - 46), xlabel, fill="#111827", font=font(27), anchor="mm")

    # Rotated y label.
    label = Image.new("RGBA", (260, 50), (255, 255, 255, 0))
    label_draw = ImageDraw.Draw(label)
    label_draw.text((130, 25), "PSM count", fill="#111827", font=font(27), anchor="mm")
    img.paste(label.rotate(90, expand=True), (25, int(margin_top + plot_h / 2 - 130)), label.rotate(90, expand=True))

    qs = pd.Series(values).quantile([0.01, 0.5, 0.99])
    stats = [
        f"n = {len(values):,}",
        f"median = {fmt_num(qs.loc[0.5])}",
        f"1%-99% = {fmt_num(qs.loc[0.01])} to {fmt_num(qs.loc[0.99])}",
    ]
    box_w, box_h = 430, 145
    box_x = margin_left + plot_w - box_w - 30
    box_y = margin_top + 30
    draw.rounded_rectangle([box_x, box_y, box_x + box_w, box_y + box_h], radius=16, fill="white", outline="#D1D5DB", width=2)
    for i, txt in enumerate(stats):
        draw.text((box_x + 26, box_y + 33 + i * 40), txt, fill="#111827", font=font(26), anchor="lm")

    img.save(path, "PNG", optimize=True)


def draw_threshold_curve(thresholds: pd.DataFrame, path: Path) -> None:
    width, height = 1800, 1050
    margin_left, margin_right, margin_top, margin_bottom = 160, 100, 130, 145
    plot_w = width - margin_left - margin_right
    plot_h = height - margin_top - margin_bottom
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)

    x_vals = thresholds["threshold"].to_numpy(float)
    fdr_pct = thresholds["fdr"].to_numpy(float) * 100.0
    kept_pct = thresholds["fraction_accepted"].to_numpy(float) * 100.0
    lo, hi = float(x_vals.min()), float(x_vals.max())
    x_pad = (hi - lo) * 0.05
    lo, hi = lo - x_pad, hi + x_pad
    y_max = max(float(fdr_pct.max()) * 1.15, 1.0)

    def xmap(x: float) -> float:
        return margin_left + (x - lo) / (hi - lo) * plot_w

    def ymap(y: float) -> float:
        return margin_top + plot_h - y / y_max * plot_h

    draw.rectangle(
        [margin_left, margin_top, margin_left + plot_w, margin_top + plot_h],
        fill="#F9FAFB",
        outline="#E5E7EB",
        width=2,
    )
    for tick in nice_ticks(0, y_max, 6):
        y = ymap(tick)
        draw.line([margin_left, y, margin_left + plot_w, y], fill="#E5E7EB", width=2)
        draw.text((margin_left - 18, y), f"{tick:.2f}%", fill="#4B5563", font=font(24), anchor="rm")

    # FDR line.
    pts = [(xmap(float(x)), ymap(float(y))) for x, y in zip(x_vals, fdr_pct)]
    if len(pts) > 1:
        draw.line(pts, fill="#7F1D1D", width=6, joint="curve")
    for x, y in pts:
        draw.ellipse([x - 8, y - 8, x + 8, y + 8], fill=DECOY, outline="white", width=2)

    # Kept percentage as muted labels along x-axis.
    for x, kept in zip(x_vals, kept_pct):
        xp = xmap(float(x))
        draw.text((xp, margin_top + plot_h - 22), f"{kept:.1f}%", fill="#1F3A5F", font=font(21), anchor="mm")

    draw.line([margin_left, margin_top + plot_h, margin_left + plot_w, margin_top + plot_h], fill="#111827", width=3)
    draw.line([margin_left, margin_top, margin_left, margin_top + plot_h], fill="#111827", width=3)
    for tick in nice_ticks(lo, hi, 7):
        x = xmap(tick)
        draw.line([x, margin_top + plot_h, x, margin_top + plot_h + 11], fill="#111827", width=2)
        draw.text((x, margin_top + plot_h + 32), fmt_num(tick), fill="#4B5563", font=font(24), anchor="mt")

    draw.text((width / 2, 58), "Estimated FDR by Neural Score Threshold", fill="#111827", font=font(44, bold=True), anchor="mm")
    draw.text((margin_left + plot_w / 2, height - 46), "Neural confidence score threshold", fill="#111827", font=font(27), anchor="mm")
    label = Image.new("RGBA", (360, 50), (255, 255, 255, 0))
    label_draw = ImageDraw.Draw(label)
    label_draw.text((180, 25), "Estimated TDA-FDR", fill="#111827", font=font(27), anchor="mm")
    rotated = label.rotate(90, expand=True)
    img.paste(rotated, (25, int(margin_top + plot_h / 2 - 180)), rotated)

    legend_x, legend_y = margin_left + plot_w - 470, margin_top + 28
    draw.rounded_rectangle([legend_x, legend_y, legend_x + 440, legend_y + 118], radius=16, fill="white", outline="#D1D5DB", width=2)
    draw.line([legend_x + 25, legend_y + 35, legend_x + 95, legend_y + 35], fill="#7F1D1D", width=6)
    draw.ellipse([legend_x + 53, legend_y + 27, legend_x + 69, legend_y + 43], fill=DECOY, outline="white", width=2)
    draw.text((legend_x + 115, legend_y + 35), "Estimated FDR", fill="#111827", font=font(25), anchor="lm")
    draw.text((legend_x + 25, legend_y + 80), "Blue labels show % of PSMs kept", fill="#1F3A5F", font=font(23), anchor="lm")

    img.save(path, "PNG", optimize=True)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "DADCE0") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def style_table(table, header: bool = True) -> None:
    table.autofit = False
    for row_i, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = TEXT
            if header and row_i == 0:
                set_cell_shading(cell, HEADER_FILL)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(17, 24, 39)


def set_col_widths(table, widths: Iterable[float]) -> None:
    widths = list(widths)
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def add_metric_table(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "Metric"
    table.rows[0].cells[1].text = "Value"
    table.rows[0].cells[2].text = "Interpretation"
    for metric, value, note in rows:
        cells = table.add_row().cells
        cells[0].text = metric
        cells[1].text = value
        cells[2].text = note
    set_col_widths(table, [1.75, 1.25, 3.35])
    style_table(table)


def add_callout(doc: Document, label: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, CALLOUT_FILL)
    set_cell_border(cell, "D9E2F3")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"{label}: ")
    run.bold = True
    run.font.color.rgb = DARK_BLUE
    p.add_run(body)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.167


def add_footer(doc: Document) -> None:
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Neural TDA-FDR report")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def add_figure(doc: Document, img_path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(6.25))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    cap_run = cap.add_run(caption)
    cap_run.italic = True
    cap_run.font.size = Pt(9)
    cap_run.font.color.rgb = MUTED


def build_report() -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    plot_dir = OUT_DIR / "plots"
    plot_dir.mkdir(exist_ok=True)

    summary = pd.read_csv(SUMMARY_CSV).iloc[0]
    thresholds = pd.read_csv(THRESHOLDS_CSV)
    psms = pd.read_csv(PSMS_CSV)
    if psms["is_decoy"].dtype != bool:
        psms["is_decoy"] = psms["is_decoy"].astype(str).str.lower().eq("true")

    psms["is_correct_target"] = (~psms["is_decoy"]) & (psms["sequence"] == psms["true_sequence"])
    psms["is_incorrect_target"] = (~psms["is_decoy"]) & (~psms["is_correct_target"])

    n_total = int(summary["n_total"])
    n_targets = int(summary["n_target_top1"])
    n_decoys = int(summary["n_decoy_top1"])
    fdr_top1 = float(summary["fdr_top1"])
    accepted_1pct = psms[(psms["qvalue"] <= 0.01) & (~psms["is_decoy"])]
    exact_correct = int(psms["is_correct_target"].sum())
    exact_accuracy = exact_correct / n_total
    exact_accepted_accuracy = int(accepted_1pct["is_correct_target"].sum()) / len(accepted_1pct)

    plots = {
        "decoy_score": plot_dir / "decoy_confidence_score_distribution.png",
        "decoy_log": plot_dir / "decoy_log_score_distribution.png",
        "target_score": plot_dir / "target_confidence_score_distribution.png",
        "target_log": plot_dir / "target_log_score_distribution.png",
        "threshold_curve": plot_dir / "threshold_fdr_curve.png",
    }
    decoy = psms[psms["is_decoy"]]
    target = psms[~psms["is_decoy"]]
    draw_histogram(
        decoy["score"].to_numpy(float),
        title="Decoy Confidence Score Distribution",
        xlabel="Confidence score = exp(mean log-probability)",
        color=DECOY,
        path=plots["decoy_score"],
        bins=80,
        xlim=(0.0, max(0.5, float(decoy["score"].max()) * 1.05)),
    )
    draw_histogram(
        decoy["mean_logp"].to_numpy(float),
        title="Decoy Log-Score Distribution",
        xlabel="Mean log-probability = log(confidence score)",
        color=DECOY,
        path=plots["decoy_log"],
        bins=80,
    )
    draw_histogram(
        target["score"].to_numpy(float),
        title="Target Confidence Score Distribution",
        xlabel="Confidence score = exp(mean log-probability)",
        color=TARGET,
        path=plots["target_score"],
        bins=100,
        xlim=(0.0, 1.0),
    )
    draw_histogram(
        target["mean_logp"].to_numpy(float),
        title="Target Log-Score Distribution",
        xlabel="Mean log-probability = log(confidence score)",
        color=TARGET,
        path=plots["target_log"],
        bins=100,
    )
    draw_threshold_curve(thresholds, plots["threshold_curve"])

    doc = Document()
    configure_doc(doc)
    add_footer(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title_run = title.add_run("Neural Target-Decoy FDR Report")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle_run = subtitle.add_run("Two-stage retrieval with InstaNovo/Casanovo-DB-style neural rescoring")
    subtitle_run.font.name = "Calibri"
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = MUTED

    add_callout(
        doc,
        "Bottom line",
        "The target-decoy competition implementation is internally consistent and gives an estimated 1.57% top-1 FDR overall. "
        "At q <= 0.01 it accepts 108,346 target PSMs at approximately 1% estimated FDR, corresponding to a neural score cutoff of about 0.056. "
        "The expanded score-threshold sweep agrees with this: score >= 0.06 keeps 97.6% of PSMs and lowers estimated FDR to 0.79%. "
        "Because benchmark labels are available, exact sequence accuracy should be reported alongside TDA-FDR: 81.98% overall and about 84.11% among q <= 0.01 accepted targets.",
    )

    add_heading(doc, "Executive Summary", 1)
    add_bullet(doc, "The code follows the standard target-decoy competition pattern: targets and reverse-inner decoys compete in the same candidate pool, and FDR is estimated from decoy winners over target winners.")
    add_bullet(doc, "The score distributions show strong target/decoy separation: target median confidence is 0.516, while decoy median confidence is 0.0598.")
    add_bullet(doc, "The headline TDA result is good for this setup: 108,346 accepted target PSMs at q <= 0.01.")
    add_bullet(doc, "A practical raw-score cutoff near the 1% FDR boundary is around 0.056-0.060; score >= 0.06 gives 0.79% estimated FDR while keeping 97.6% of PSMs.")
    add_bullet(doc, "The main caveat is interpretation: exact benchmark accuracy is lower than the TDA estimate because many mistakes are target-vs-target confusions, especially similar modified and unmodified peptide variants.")

    add_heading(doc, "Key Results", 1)
    add_metric_table(
        doc,
        [
            ("Spectra searched", f"{n_total:,}", "Input spectra after excluding targets without valid 1:1 decoys."),
            ("Unique targets", f"{int(summary['n_unique_targets']):,}", "Unique modified peptide sequences before decoy filtering."),
            ("Targets with decoys", f"{int(summary['n_targets_with_decoys']):,}", "Strict 1:1 target-decoy pairs used in the competition."),
            ("Top-1 target winners", f"{n_targets:,}", "Final neural winner was a target peptide."),
            ("Top-1 decoy winners", f"{n_decoys:,}", "Final neural winner was a decoy peptide."),
            ("Estimated top-1 FDR", f"{fdr_top1:.2%}", "Computed as decoy winners divided by target winners."),
            ("q <= 0.01 accepted targets", f"{len(accepted_1pct):,}", "Accepted target-only PSM set at estimated 1% FDR."),
            ("q <= 0.01 score cutoff", f"{accepted_1pct['score'].min():.4f}", "Minimum neural score among accepted q<=1% target PSMs."),
            ("score >= 0.06 FDR", "0.79%", "Mild raw-score cutoff; keeps 97.6% of PSMs."),
            ("Exact top-1 accuracy", f"{exact_accuracy:.2%}", "Benchmark-only exact match against true_sequence."),
            ("Exact accuracy within q <= 0.01 targets", f"{exact_accepted_accuracy:.2%}", "Benchmark-only exact match after the 1% q-value filter."),
        ],
    )

    add_heading(doc, "Threshold Behavior", 1)
    doc.add_paragraph(
        "The expanded raw-score sweep is now informative. Weak filtering near 0.03-0.04 keeps almost the entire dataset and remains close to the overall 1.57% FDR. "
        "The useful transition occurs around 0.05-0.06: score >= 0.05 gives 1.32% estimated FDR, while score >= 0.06 gives 0.79% estimated FDR. "
        "This aligns with the q <= 0.01 accepted set, whose implied score cutoff is about 0.056."
    )
    table = doc.add_table(rows=1, cols=6)
    for i, header in enumerate(["Threshold", "Accepted", "Targets", "Decoys", "Estimated FDR", "% kept"]):
        table.rows[0].cells[i].text = header
    for _, row in thresholds.iterrows():
        cells = table.add_row().cells
        cells[0].text = f"{row['threshold']:.3g}"
        cells[1].text = f"{int(row['n_accepted']):,}"
        cells[2].text = f"{int(row['target_hits']):,}"
        cells[3].text = f"{int(row['decoy_hits']):,}"
        cells[4].text = f"{float(row['fdr']):.2%}"
        cells[5].text = f"{float(row['fraction_accepted']):.1%}"
    set_col_widths(table, [0.85, 1.15, 1.05, 0.9, 1.15, 0.9])
    style_table(table)
    add_figure(doc, plots["threshold_curve"], "Figure 1. Estimated FDR decreases sharply once the neural score threshold moves through the target/decoy overlap region.")

    add_heading(doc, "Interpretation for Discussion", 1)
    doc.add_paragraph(
        "The clean way to present the result is to separate the target-decoy estimate from benchmark exact-match accuracy. "
        "TDA-FDR answers: how many accepted target PSMs are expected to be false under the decoy model? "
        "Exact-match accuracy answers: how often did the selected target sequence equal the benchmark label?"
    )
    add_bullet(doc, "Supervisor-facing headline: the neural TDA-FDR machinery is working and gives approximately 1% estimated FDR at q <= 0.01.")
    add_bullet(doc, "Required caveat: this does not mean 99% exact peptide correctness on the benchmark; exact sequence accuracy is about 84% within the accepted q <= 0.01 target set.")
    add_bullet(doc, "Likely cause: target-vs-target errors, including modification-state or modification-localization confusions, are not fully captured by the decoy count.")
    add_bullet(doc, "Recommended next step: report both TDA-FDR and exact benchmark accuracy, then inspect high-confidence incorrect target calls to understand modification-related failure modes.")

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "Score Distribution Plots", 1)
    doc.add_paragraph(
        "The figures below show the raw confidence score and log-score distributions for target and decoy winners. "
        "The strong separation between targets and decoys supports the TDA estimate, while the benchmark accuracy caveat remains important for exact-sequence interpretation."
    )

    add_heading(doc, "Decoy Distributions", 2)
    add_figure(doc, plots["decoy_score"], "Figure 2. Decoy confidence scores are concentrated at low values.")
    add_figure(doc, plots["decoy_log"], "Figure 3. Decoy mean log-probabilities cluster far below the target median.")

    add_heading(doc, "Target Distributions", 2)
    add_figure(doc, plots["target_score"], "Figure 4. Target confidence scores are broadly shifted toward higher values.")
    add_figure(doc, plots["target_log"], "Figure 5. Target mean log-probabilities show the corresponding log-scale separation.")

    add_heading(doc, "Suggested Wording", 1)
    p = doc.add_paragraph()
    p.add_run("Suggested supervisor update: ").bold = True
    p.add_run(
        "We implemented neural target-decoy competition after two-stage retrieval and InstaNovo/Casanovo-DB-style rescoring. "
        "The TDC procedure is internally consistent: with a 1:1 reverse-inner decoy database, the overall estimated top-1 FDR is 1.57%, "
        "and applying q <= 0.01 yields 108,346 accepted target PSMs at approximately 1% estimated FDR. "
        "The expanded score-threshold sweep confirms that the practical 1% boundary is near score 0.056-0.060; at score >= 0.06, the estimated FDR is 0.79% while keeping 97.6% of PSMs. "
        "The target and decoy score distributions are well separated. "
        "However, because benchmark labels are available, we also report exact sequence accuracy: about 82% overall and 84% among q <= 0.01 accepted targets. "
        "This indicates residual target-vs-target errors, especially likely around modified or closely related peptide variants, so the FDR result is promising but should be presented together with benchmark exact-match accuracy."
    )

    output_path = OUT_DIR / "neural_tda_fdr_report.docx"
    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    print(build_report())
