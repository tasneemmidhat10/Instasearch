"""Generate an editable DOCX plan for Target-Decoy Competition (TDC) at 1% FDR
using the neural rescoring score (geometric-mean log-probability exponentiated).
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def shade_row(row, color_hex):
    for cell in row.cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), color_hex)
        tc_pr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.75 + 0.75 * level)
    return p


def add_numbered(doc, text):
    return doc.add_paragraph(text, style="List Number")


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    # light grey shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    return p


def main():
    doc = Document()

    # Default body style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ----------------------------------------------------------------------
    # Title block
    # ----------------------------------------------------------------------
    title = doc.add_heading(
        "Target–Decoy Competition Analysis at 1% FDR using the Neural Rescoring Score",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run(
        "Comparative evaluation against Gabrielle’s pipeline on the Nine-Species benchmark "
        "with both global and local FDR control"
    )
    sub_run.italic = True
    sub_run.font.size = Pt(12)

    meta = doc.add_paragraph()
    meta.add_run("Author: ").bold = True
    meta.add_run("Tasneem  •  ")
    meta.add_run("Date: ").bold = True
    meta.add_run("2026-05-20  •  ")
    meta.add_run("Status: ").bold = True
    meta.add_run("Draft for review")

    doc.add_paragraph()

    # ----------------------------------------------------------------------
    # 1. Motivation
    # ----------------------------------------------------------------------
    add_heading(doc, "1. Motivation and Objective", level=1)
    add_para(
        doc,
        "De novo peptide identifications from the dual-encoder retrieval-and-rescoring "
        "pipeline are currently ranked by cosine similarity between the spectrum and "
        "peptide embeddings. While cosine similarity is a useful retrieval signal, it is "
        "not well-calibrated as a confidence score: its distribution depends strongly "
        "on charge state, peptide length, and the local geometry of the embedding "
        "space, which makes a single cosine threshold a poor proxy for the true error "
        "rate of the returned identifications.",
    )
    add_para(
        doc,
        "We therefore propose to perform a Target–Decoy Competition (TDC) analysis on "
        "the retrieval output and to control the False Discovery Rate (FDR) at 1%, "
        "using the neural rescoring score — defined as the exponential of the "
        "log-probability geometric mean produced by the rescoring model — as the "
        "discriminating statistic. The same protocol will be executed against "
        "Gabrielle’s pipeline and against our test set (the Nine-Species benchmark) "
        "so that the two systems can be compared head-to-head at a matched, "
        "calibrated confidence level.",
    )

    add_heading(doc, "1.1 Why the rescoring score and not cosine?", level=2)
    add_bullet(
        doc,
        "Cosine is a similarity in embedding space; its scale shifts with charge, "
        "length, and dataset, so an FDR threshold learnt on one slice does not "
        "transfer to another.",
    )
    add_bullet(
        doc,
        "The neural rescoring score s = exp((1/L) · Σ log p(aa_i | spectrum)) is the "
        "geometric mean of the model’s per-residue probabilities. It is length-normalised "
        "by construction and lives on [0, 1].",
    )
    add_bullet(
        doc,
        "Because it is derived from a likelihood, it behaves much more like a "
        "probabilistic confidence than cosine and is a more appropriate input to the "
        "decoy-competition machinery.",
    )

    # ----------------------------------------------------------------------
    # 2. Scope
    # ----------------------------------------------------------------------
    add_heading(doc, "2. Scope and Deliverables", level=1)
    add_bullet(doc, "Compute global FDR and local FDR (q-values and posterior error probabilities) "
                    "for every PSM returned by our pipeline.")
    add_bullet(doc, "Reproduce the same TDC protocol on Gabrielle’s pipeline using its own "
                    "rescoring output, so both pipelines are scored on a like-for-like basis.")
    add_bullet(doc, "Run the analysis on the Nine-Species benchmark (leave-one-species-out, "
                    "as is standard for that benchmark).")
    add_bullet(doc, "Report number of accepted PSMs at 1% global FDR per species and overall, "
                    "with calibration and sensitivity plots.")
    add_bullet(doc, "Produce a reusable script and notebook so the analysis can be re-run "
                    "after any model retraining.")

    # ----------------------------------------------------------------------
    # 3. Definitions
    # ----------------------------------------------------------------------
    add_heading(doc, "3. Definitions and Notation", level=1)
    add_para(doc, "Let a PSM be a (spectrum, candidate peptide) pair. For each PSM we record:")
    add_bullet(doc, "label ∈ {target, decoy}")
    add_bullet(doc, "rescoring score s = exp((1/L) · Σ_{i=1..L} log p(aa_i | spectrum, context))")
    add_bullet(doc, "any auxiliary features used by the rescorer (kept for diagnostics only).")

    add_para(doc, "Global FDR at threshold t (target-decoy estimator):", italic=True)
    add_code(doc, "FDR_global(t) = (#decoys with s ≥ t + 1) / max(1, #targets with s ≥ t)")
    add_para(
        doc,
        "The +1 correction (Storey-style) gives a conservative estimate when decoy "
        "counts are small. The q-value of a PSM is the minimum FDR at which that PSM "
        "would be accepted:",
        italic=True,
    )
    add_code(doc, "q(s_i) = min_{t ≤ s_i} FDR_global(t)")

    add_para(doc, "Local FDR (posterior error probability, PEP) at score s:", italic=True)
    add_code(
        doc,
        "PEP(s) = P(decoy | s) = π0 · f_decoy(s) / ( π0 · f_decoy(s) + (1 − π0) · f_target(s) )",
    )
    add_para(
        doc,
        "where f_target and f_decoy are smoothed score densities on the target and decoy "
        "subsets respectively, and π0 is the estimated null proportion among targets "
        "(estimated, e.g., with Storey’s method on the score histogram).",
        italic=True,
    )

    # ----------------------------------------------------------------------
    # 4. Decoy generation
    # ----------------------------------------------------------------------
    add_heading(doc, "4. Decoy Generation Strategy", level=1)
    add_para(
        doc,
        "Decoys must be (a) indistinguishable from targets in any property the model "
        "could exploit a priori (composition, length distribution, mass), and "
        "(b) genuinely null with respect to the spectra. We will use shuffled-target "
        "decoys as the primary strategy and reversed-target decoys as a robustness check.",
    )
    add_bullet(doc, "Primary: shuffle the internal residues of each target peptide while "
                    "keeping termini fixed; reject shuffles that collide with any target.")
    add_bullet(doc, "Secondary (sanity): reverse the target sequence keeping termini fixed.")
    add_bullet(doc, "Decoys inherit the precursor m/z window of their target; the rescoring "
                    "model sees them with the same input format as targets.")
    add_bullet(doc, "1:1 target:decoy ratio. Decoys are scored by the same rescoring "
                    "model under identical conditions to targets.")

    add_para(
        doc,
        "Note: with both pipelines we must be careful that decoys are generated "
        "consistently. For Gabrielle’s pipeline we will use her decoy-generation routine "
        "if available, and otherwise apply the same shuffling procedure to her candidate "
        "sequences before passing them through her rescorer.",
        italic=True,
    )

    # ----------------------------------------------------------------------
    # 5. Datasets and pipelines
    # ----------------------------------------------------------------------
    add_heading(doc, "5. Datasets and Pipelines", level=1)

    add_heading(doc, "5.1 Test set — Nine-Species benchmark", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Species"
    hdr[1].text = "Role"
    hdr[2].text = "Notes / placeholder"
    shade_row(table.rows[0], "D9E2F3")
    species = [
        ("Mus musculus", "held-out", ""),
        ("Homo sapiens", "held-out", ""),
        ("Saccharomyces cerevisiae", "held-out", ""),
        ("Methanosarcina mazei", "held-out", ""),
        ("Vigna mungo", "held-out", ""),
        ("Candidatus endoloripes", "held-out", ""),
        ("Bacillus subtilis", "held-out", ""),
        ("Solanum lycopersicum", "held-out", ""),
        ("Apis mellifera", "held-out", ""),
    ]
    for name, role, note in species:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = role
        row[2].text = note

    add_heading(doc, "5.2 Pipelines under comparison", level=2)
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = "Light Grid Accent 1"
    hdr = table2.rows[0].cells
    hdr[0].text = "Pipeline"
    hdr[1].text = "Retrieval"
    hdr[2].text = "Rescoring / score used"
    shade_row(table2.rows[0], "D9E2F3")
    rows = [
        ("Ours", "Dual-encoder + HNSW (FAISS)",
         "Neural rescoring: s = exp(mean log p)"),
        ("Gabrielle’s", "<fill in>", "Her rescorer’s analogous neural score "
                                     "(exp mean log p); confirm exact definition."),
    ]
    for r in rows:
        cells = table2.add_row().cells
        for i, val in enumerate(r):
            cells[i].text = val

    # ----------------------------------------------------------------------
    # 6. Protocol
    # ----------------------------------------------------------------------
    add_heading(doc, "6. Analysis Protocol", level=1)

    add_heading(doc, "6.1 Per-pipeline steps", level=2)
    add_numbered(doc, "Generate decoy peptide set (shuffled targets, 1:1 ratio).")
    add_numbered(doc, "Score all target and decoy candidates with the pipeline’s neural rescorer; "
                      "record s = exp(mean log p) and metadata (charge, length, species).")
    add_numbered(doc, "Per spectrum, perform target–decoy competition: keep the single "
                      "best-scoring candidate (target or decoy). This collapses each spectrum "
                      "to one PSM and enforces the competition.")
    add_numbered(doc, "Sort competition winners by s in descending order.")
    add_numbered(doc, "Compute global FDR and q-values as in §3.")
    add_numbered(doc, "Fit f_target, f_decoy (kernel density or monotone spline) and "
                      "estimate π0 to get local FDR / PEP per PSM.")
    add_numbered(doc, "Threshold at 1% global FDR; report the accepted set.")

    add_heading(doc, "6.2 Cross-pipeline comparison", level=2)
    add_bullet(doc, "Apply 6.1 independently to our pipeline and to Gabrielle’s, on the same "
                    "spectra from each Nine-Species held-out species.")
    add_bullet(doc, "Report, per species and overall: #accepted PSMs at 1% FDR, peptide-level "
                    "accepted counts (after grouping PSMs by peptide), and amino-acid-level "
                    "accuracy on accepted PSMs (using the species ground-truth peptide).")
    add_bullet(doc, "Use the same decoys for both pipelines wherever possible.")

    # ----------------------------------------------------------------------
    # 7. Diagnostics and figures
    # ----------------------------------------------------------------------
    add_heading(doc, "7. Diagnostics and Required Figures", level=1)
    add_bullet(doc, "Score histograms of targets vs decoys per pipeline and per species.")
    add_bullet(doc, "Calibration plot: empirical FDR (using ground-truth species labels) vs "
                    "TDC-estimated FDR. The closer to y=x, the better calibrated.")
    add_bullet(doc, "q-value vs number-of-accepted-PSMs curves (one line per pipeline).")
    add_bullet(doc, "PEP (local FDR) distribution among accepted PSMs.")
    add_bullet(doc, "Per-species bar chart: accepted PSMs at 1% FDR — ours vs Gabrielle’s.")

    # ----------------------------------------------------------------------
    # 8. Validation
    # ----------------------------------------------------------------------
    add_heading(doc, "8. Validation Checks", level=1)
    add_bullet(doc, "Decoy sanity: target and decoy score distributions should overlap in the "
                    "low-score region but diverge in the high-score tail. If they are "
                    "indistinguishable, the rescorer adds no information; if they are "
                    "separated everywhere, decoys are too easy and FDR is under-estimated.")
    add_bullet(doc, "FDR honesty: with shuffled decoys, the empirical FDR (computed against "
                    "Nine-Species ground truth) at the 1% TDC threshold should be ≤ ~1% "
                    "(small over-shoot acceptable; large over-shoot indicates miscalibration).")
    add_bullet(doc, "Stability: rerun with reversed decoys; conclusions should not change "
                    "qualitatively.")
    add_bullet(doc, "Permutation test: shuffle the target/decoy labels of all PSMs and verify "
                    "that the number of accepted PSMs at 1% FDR drops to near zero.")

    # ----------------------------------------------------------------------
    # 9. Implementation notes
    # ----------------------------------------------------------------------
    add_heading(doc, "9. Implementation Notes", level=1)
    add_bullet(doc, "Add a TDC module under src/retrieval/ (e.g. src/retrieval/tdc.py) that "
                    "takes a dataframe with columns [spectrum_id, peptide, label, score] and "
                    "returns q-values and PEPs.")
    add_bullet(doc, "Reuse the existing scoring path so the s = exp(mean log p) value is "
                    "logged alongside cosine — do not replace cosine, log both.")
    add_bullet(doc, "Persist decoy sequences with deterministic seeds so runs are reproducible.")
    add_bullet(doc, "Notebook: notebooks/tdc_fdr_analysis.ipynb — load PSM tables for both "
                    "pipelines, run TDC, render the figures from §7.")

    add_para(doc, "Suggested function signatures:", italic=True)
    add_code(
        doc,
        "def compute_qvalues(scores: np.ndarray, is_decoy: np.ndarray) -> np.ndarray: ...\n"
        "def compute_pep(scores: np.ndarray, is_decoy: np.ndarray,\n"
        "                pi0: float | None = None) -> np.ndarray: ...\n"
        "def tdc_compete(psms: pd.DataFrame, score_col: str = 'neural_score',\n"
        "                spectrum_col: str = 'spectrum_id') -> pd.DataFrame: ...",
    )

    # ----------------------------------------------------------------------
    # 10. Risks
    # ----------------------------------------------------------------------
    add_heading(doc, "10. Risks and Mitigations", level=1)
    table3 = doc.add_table(rows=1, cols=2)
    table3.style = "Light Grid Accent 1"
    table3.rows[0].cells[0].text = "Risk"
    table3.rows[0].cells[1].text = "Mitigation"
    shade_row(table3.rows[0], "D9E2F3")
    risks = [
        ("Decoys too easy to discriminate → FDR under-estimated.",
         "Use shuffle-with-termini-fixed; check score overlap; cross-check with reversed decoys."),
        ("Score not actually well-calibrated either (just better than cosine).",
         "Quantify with empirical-vs-estimated FDR plot; if needed, post-hoc calibrate "
         "with Percolator-style SVM or isotonic regression on a held-out species."),
        ("Gabrielle’s pipeline uses a different score definition.",
         "Confirm exact formula with her; if different, use her own analogous "
         "‘confidence’ output and document the asymmetry."),
        ("Per-species sample sizes are small → noisy q-values.",
         "Report overall pooled FDR as the headline; per-species numbers as secondary."),
    ]
    for r in risks:
        cells = table3.add_row().cells
        cells[0].text = r[0]
        cells[1].text = r[1]

    # ----------------------------------------------------------------------
    # 11. Timeline
    # ----------------------------------------------------------------------
    add_heading(doc, "11. Timeline (placeholder — edit as needed)", level=1)
    table4 = doc.add_table(rows=1, cols=3)
    table4.style = "Light Grid Accent 1"
    table4.rows[0].cells[0].text = "Phase"
    table4.rows[0].cells[1].text = "Deliverable"
    table4.rows[0].cells[2].text = "Target date"
    shade_row(table4.rows[0], "D9E2F3")
    phases = [
        ("Decoy generator + score export", "Decoys + PSM tables for both pipelines", "<edit>"),
        ("TDC module + unit tests", "src/retrieval/tdc.py with q-value / PEP", "<edit>"),
        ("Run on Nine-Species", "Per-species CSVs of accepted PSMs at 1% FDR", "<edit>"),
        ("Figures + write-up", "Notebook with calibration & comparison plots", "<edit>"),
    ]
    for r in phases:
        cells = table4.add_row().cells
        for i, val in enumerate(r):
            cells[i].text = val

    # ----------------------------------------------------------------------
    # 12. Open questions
    # ----------------------------------------------------------------------
    add_heading(doc, "12. Open Questions (for review)", level=1)
    add_bullet(doc, "Should the competition be at the spectrum level (one PSM per spectrum) or "
                    "also at the peptide level (further collapse across charges/scans)?")
    add_bullet(doc, "Do we want a single pooled FDR across all nine species, or species-by-species "
                    "FDR control (the latter is stricter)?")
    add_bullet(doc, "Is 1:1 target:decoy enough, or should we use k>1 decoys per target and apply "
                    "the corresponding TDC correction?")
    add_bullet(doc, "Do we keep cosine alongside as a baseline FDR analysis to show the "
                    "calibration gap?")

    # ----------------------------------------------------------------------
    # Save
    # ----------------------------------------------------------------------
    out_path = r"c:\Users\tasne\Desktop\project\docs\TDC_FDR_Analysis_Plan.docx"
    doc.save(out_path)
    print(f"WROTE {out_path}")


if __name__ == "__main__":
    main()
