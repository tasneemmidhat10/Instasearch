"""Generate the revised post-neural-rescoring TDC/FDR analysis plan."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_PATH = Path(r"C:\Users\tasne\Desktop\project\docs\TDC_FDR_Analysis_Plan_Post_Neural_Rescoring.docx")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_paragraph(doc: Document, text: str, style: str | None = None):
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p


def add_bullet(doc: Document, text: str):
    return add_paragraph(doc, text, "List Bullet")


def add_number(doc: Document, text: str):
    return add_paragraph(doc, text, "List Number")


def add_code(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F6F9")
    p_pr.append(shd)
    return p


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.width = Inches(widths[i])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "E8EEF5")
        set_cell_margins(cell)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            cells[i].width = Inches(widths[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[i])
    return table


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 14, 6),
        ("Heading 2", 13, "2E74B5", 10, 4),
        ("Heading 3", 12, "1F4D78", 8, 3),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for name in ["List Bullet", "List Number"]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15


def main() -> None:
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Target-Decoy Competition at 1% FDR After Neural Rescoring")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string("0B2545")
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.add_run(
        "Revised implementation plan for the dual-encoder retrieval pipeline and Nine-Species comparison"
    ).italic = True

    doc.add_heading("1. Objective", level=1)
    add_paragraph(
        doc,
        "Run Target-Decoy Competition only after the neural rescoring phase. Cosine similarity remains "
        "a stage-1 retrieval signal for generating candidates, but it is not used as the final FDR "
        "confidence statistic. The final competition score is the neural rescoring score:",
    )
    add_code(doc, "s = exp((1 / L) * sum_i log p(aa_i | spectrum, peptide prefix))")
    add_paragraph(
        doc,
        "This is the geometric mean of the model's per-residue probabilities. It is length-normalized, "
        "bounded in [0, 1], and is the score used for target-decoy ordering, q-values, local FDR/PEP, "
        "and the 1% accepted set.",
    )

    doc.add_heading("2. Required Pipeline Order", level=1)
    for step in [
        "Encode spectra and peptide targets with the dual encoder.",
        "Deduplicate target peptides by modified_sequence.",
        "Generate one valid reverse-inner decoy per usable unique target peptide.",
        "Search each spectrum against the combined target plus decoy database using cosine top-k.",
        "Neurally rescore every retrieved target and decoy candidate with the InstaNovo/Casanovo-DB mean log-probability scorer.",
        "For each spectrum, keep the single best neural-score candidate, target or decoy.",
        "Sort these competition winners by s descending.",
        "Compute global TDC FDR, monotone q-values, local FDR/PEP, and the accepted target PSMs at q <= 0.01.",
    ]:
        add_number(doc, step)

    doc.add_heading("3. Decoy Strategy", level=1)
    add_paragraph(
        doc,
        "Use reverse-inner decoys as the implemented primary decoy strategy. The first and last residue "
        "tokens are preserved and the internal tokens are reversed. A decoy is rejected if it is identical "
        "to its target, collides with any target sequence, or duplicates a previous decoy.",
    )
    add_paragraph(
        doc,
        "To preserve a strict 1:1 target-decoy database, targets that cannot produce a valid decoy are "
        "excluded from the neural TDC analysis, and spectra whose ground-truth target is excluded are also "
        "excluded from the competition. The result records n_input_spectra and n_excluded_no_decoy so the "
        "filter is visible.",
    )

    doc.add_heading("4. FDR and PEP Definitions", level=1)
    add_code(
        doc,
        "FDR(t) = min(1, decoy_winners_with_score >= t / max(1, target_winners_with_score >= t))\n"
        "qvalue(i) = min FDR(t) over thresholds t <= score_i\n"
        "PEP(s) = pi0 * f_decoy(s) / (pi0 * f_decoy(s) + (1 - pi0) * f_target(s))",
    )
    add_paragraph(
        doc,
        "The implementation fits target and decoy score densities using a Gaussian KDE and uses the "
        "observed decoy fraction among finite-scored competition winners as the plug-in pi0 estimate. "
        "The accepted set contains target winners only, with qvalue <= 0.01.",
    )

    doc.add_heading("5. Code Contract", level=1)
    add_table(
        doc,
        ["Function", "Responsibility", "Plan requirement satisfied"],
        [
            [
                "src.retrieval.rerank.compute_neural_tda_fdr",
                "Build target plus reverse-inner decoy DB, retrieve cosine top-k, rescore candidates neurally, and keep the best target/decoy winner per spectrum.",
                "Post-neural-rescoring target-decoy competition.",
            ],
            [
                "src.retrieval.search.compute_tda_fdr",
                "Public FDR entry point. When rescorer_model is supplied, delegates to neural TDC and adds PEP, pi0, accepted_psms, and diagnostics.",
                "Global FDR, q-values, local FDR/PEP, and accepted 1% target PSMs.",
            ],
            [
                "src.retrieval.benchmarks.run_fdr_benchmark",
                "Benchmark wrapper that runs neural TDC when a rescorer is provided and skips legacy external-decoy cosine TDA in that mode.",
                "Ensures FDR happens after neural rescoring in benchmark runs.",
            ],
        ],
        [2.0, 2.4, 2.1],
    )

    doc.add_heading("6. Required Outputs", level=1)
    for item in [
        "competition_psms: one row per spectrum-level competition winner with sequence, is_decoy, score, mean_logp, stage1_cosine, qvalue, pep, charge, peptide_length, and accepted_1pct_fdr.",
        "accepted_psms: target-only subset of competition_psms where qvalue <= 0.01.",
        "diagnostics: target_scores, decoy_scores, sorted_scores, sorted_qvalues, accepted PEP values, and score cutoff.",
        "Run summary: n_input_spectra, n_total analyzed spectra, n_excluded_no_decoy, n_targets_with_decoys, n_target_top1, n_decoy_top1, and fdr_top1.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("7. Required Figures", level=1)
    for item in [
        "Target vs decoy score histograms using the final neural score s.",
        "q-value vs number of accepted target PSMs.",
        "PEP distribution among accepted target PSMs.",
        "Calibration curve comparing empirical error against TDC-estimated FDR when ground truth is available.",
        "Per-species accepted PSM counts at 1% FDR for our pipeline and Gabriel's pipeline.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("8. Usage Pattern", level=1)
    add_code(
        doc,
        "from src.retrieval.search import compute_tda_fdr\n\n"
        "tda = compute_tda_fdr(\n"
        "    model_spec=model_spec,\n"
        "    model_pep=model_pep,\n"
        "    rescorer_model=instanovo_model,\n"
        "    peptide_residue_set=residue_set,\n"
        "    loader=test_loader,\n"
        "    device=device,\n"
        "    modified_sequences=test_df['modified_sequence'].tolist(),\n"
        "    stage1_top_k=100,\n"
        "    score_mode='geometric_mean',\n"
        "    fdr_cutoff=0.01,\n"
        ")\n\n"
        "accepted = tda['accepted_psms']\n"
        "diagnostics = tda['diagnostics']",
    )

    doc.add_heading("9. Validation Checks", level=1)
    for item in [
        "Verify that final_scores, qvalues, pep, and accepted_psms are based on neural scores, not cosine.",
        "Check that target and decoy score distributions overlap in the low-score region.",
        "Confirm accepted_psms contains no decoy winners.",
        "Report excluded no-decoy spectra; if this number is high, inspect decoy generation failures.",
        "Compare estimated FDR against empirical error on Nine-Species ground truth.",
    ]:
        add_bullet(doc, item)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
